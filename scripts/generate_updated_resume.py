from __future__ import annotations

from pathlib import Path
import shutil
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
PUBLIC_RESUME = ROOT / "portfolio_site" / "public" / "resume.pdf"
DOCX_PATH = DOCS_DIR / "Eslam_Tarek_Nabawy_Resume_ATS_Optimized.docx"
PDF_PATH = DOCS_DIR / "Eslam_Tarek_Nabawy_Resume_ATS_Optimized.pdf"
LEGACY_DOCX_PATH = DOCS_DIR / "Eslam_Tarek_Nabawy_Resume_Updated.docx"
LEGACY_PDF_PATH = DOCS_DIR / "Eslam_Tarek_Nabawy_Resume_Updated.pdf"


RESUME = {
    "name": "Eslam Tarek Nabawy",
    "title": "Flutter Software Engineer | Mobile & Desktop Apps | Firebase | WebRTC | AI Automation",
    "contact": [
        "6th October City, Egypt",
        "+201015683693",
        "eslamtarek.dev@gmail.com",
        "linkedin.com/in/eslam-tarek-nabawy",
        "eslamnabawy.github.io/Nabawy-s-Portfolio-/",
        "github.com/EslamNabawy",
    ],
    "summary": (
        "Flutter Software Engineer and Computer Science graduate with 3+ years of hands-on freelance "
        "and project experience building cross-platform mobile and desktop applications. Strong in Dart, "
        "Flutter, Firebase, REST APIs, WebRTC, Riverpod, BLoC/Cubit, Drift, Hive, Git, and n8n/LLM "
        "automation. Built real products across Android, Windows, Firebase-backed workflows, local-first "
        "desktop systems, and private peer-to-peer communication."
    ),
    "skills": [
        ("Programming Languages", "Dart, JavaScript, Python, C++, Java, HTML, CSS"),
        ("Mobile & Frontend", "Flutter, Flutter Desktop, Android, Windows desktop, responsive UI, Figma, React.js familiarity"),
        ("State Management & Architecture", "Riverpod, BLoC, Cubit, MVC, GoRouter, repository pattern, feature-based architecture"),
        ("Backend, APIs & Data", "Firebase Auth, Firebase Realtime Database, Cloud Firestore, Firebase Storage, REST API integration, Drift, Hive, SQL, NoSQL"),
        ("Realtime & Platform", "WebRTC data channels, voice calls, video calls, file transfer flows, local persistence, release artifacts"),
        ("AI & Automation", "n8n workflow automation, OpenAI API, Claude API, LLM API integration, webhooks, AI-assisted app creation"),
        ("Tools & Engineering", "Git, GitHub, GitHub Actions, CI/CD concepts, unit testing, debugging, OOP, design patterns, data structures, algorithms, Docker familiarity, AWS familiarity"),
    ],
    "projects": [
        {
            "heading": "Brox | Offline-First Workspace Operations App | Flutter Desktop, Hive, Cubit/Bloc",
            "bullets": [
                "Built a local-first Flutter desktop operations system for coworking/front-desk workflows including check-in, checkout, live sessions, customers, rooms, desks, reservations, kitchen inventory, finance, staff, memberships, analytics, roles, audit logs, export, restore, and app health.",
                "Designed the app around GoRouter shell routing, Cubit/Bloc state management, Hive CE local persistence, JSON-backed repositories, EasyLocalization, shared UI primitives, Windows release scripts, and GitHub Actions.",
                "Maintained production-style project structure with 531 Dart files under lib/src, 98 Dart test files, 520 widgets, local schema version 3, backup/restore flows, permissions, diagnostics, and release documentation.",
            ],
        },
        {
            "heading": "Rain | Private P2P Messenger | Flutter, WebRTC, Firebase, Drift",
            "bullets": [
                "Built a private peer-to-peer chat app for Android and Windows with accepted friendships, WebRTC data-channel messaging, connection diagnostics, one-to-one file transfer, voice calls, and video calls.",
                "Separated Flutter UI, Riverpod state, Drift local persistence, Firebase Auth/Realtime Database signaling, Remote Config, session policy, and WebRTC media/data primitives across app, core, protocol, and peer packages.",
                "Implemented username sign-in, friend search, friend requests, blocking, call/file conflict guards, direct/relay route visibility, recovery states, and Android/Windows release artifacts.",
            ],
        },
        {
            "heading": "CampusSuit | Campus Life Management App | Graduation Project | Flutter, Firebase",
            "bullets": [
                "Built a Flutter campus app covering onboarding, campus activities, events, profile management, live data, Firebase-backed workflows, and REST API integrations.",
                "Structured feature screens, state management, and data-backed UI flows so graduation-project requirements could grow without turning the UI layer into backend glue.",
            ],
        },
        {
            "heading": "AI Workflow Automation | n8n, OpenAI API, Claude API, Webhooks",
            "bullets": [
                "Built n8n workflows that connect OpenAI/Claude APIs, webhooks, notifications, and external services into repeatable AI automation pipelines.",
                "Automated data processing, AI-generated content flows, smart notifications, conditional routing, and error-aware workflow paths.",
            ],
        },
        {
            "heading": "So She Picks | E-Commerce Ordering App | Flutter, Firebase",
            "bullets": [
                "Delivered responsive ordering UI, Firebase-backed order management, checkout flow, and loyalty points behavior for a mobile commerce experience.",
            ],
        },
    ],
    "experience": [
        {
            "heading": "Freelance Software Developer | Self-Employed | 6th October City, Egypt | Nov 2022 - Present",
            "bullets": [
                "Built Flutter mobile and desktop applications using Dart, Firebase, REST APIs, responsive UI, local persistence, and feature-based architecture across client and personal products.",
                "Integrated Firebase Auth, Firestore, Realtime Database, Firebase Storage, REST APIs, routing, and state management with BLoC/Cubit and Riverpod-style patterns.",
                "Delivered product workflows from requirements to release artifacts using Git/GitHub, documentation, debugging, and AI-assisted development while keeping implementation decisions explicit and reviewable.",
                "Created n8n and LLM API workflows using OpenAI, Claude, webhooks, conditional routing, and notifications to automate repetitive product and content operations.",
            ],
        },
        {
            "heading": "Programming Instructor | Local School & Private Sessions | 6th October City, Egypt | Aug 2023 - Present",
            "bullets": [
                "Taught programming fundamentals, OOP, algorithms, data structures, debugging, and practical software development through structured exercises.",
                "Mentored students through problem decomposition, implementation habits, and code-review thinking.",
            ],
        },
    ],
    "leadership": [
        {
            "heading": "Flutter Head | Google Developer Student Club | Oct 2022 - Oct 2023",
            "bullets": [
                "Led Flutter learning sessions, mentored beginner developers, and helped organize community technical events.",
                "Coordinated NextGen and Google I/O extended events with the GDSC AEA team.",
            ],
        },
        {
            "heading": "ICPC Core Team Member | Competitive Programming | Jun 2020 - Oct 2023",
            "bullets": [
                "Organized problem-solving contests and supported new members in algorithmic thinking and contest preparation.",
                "Built a three-year habit of solving difficult problems under time pressure and explaining solutions clearly.",
            ],
        },
    ],
    "education": "Bachelor of Science in Computer Science | AEA Academy, 6th October City, Egypt | Oct 2020 - Jul 2024",
    "languages": "Arabic: Native | English: Upper-Intermediate / Professional proficiency",
}


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.font.bold = True

    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(5)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(4)
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)


def add_docx_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.add_run(text.upper())


def add_docx_bullets(doc: Document, bullets: Iterable[str]) -> None:
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        run = p.add_run(bullet)
        set_run_font(run, size=10.2)


def build_docx() -> None:
    doc = Document()
    configure_docx_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run(RESUME["name"].upper())
    set_run_font(run, size=22, color=RGBColor(0, 0, 0), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    run = subtitle.add_run(RESUME["title"])
    set_run_font(run, size=11.5, color=RGBColor(70, 70, 70), bold=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(9)
    run = contact.add_run(" | ".join(RESUME["contact"]))
    set_run_font(run, size=9.2, color=RGBColor(85, 85, 85))

    add_docx_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_run_font(p.add_run(RESUME["summary"]), size=10.5)

    add_docx_section_heading(doc, "Technical Skills")
    for label, value in RESUME["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(f"{label}: "), size=10.2, bold=True)
        set_run_font(p.add_run(value), size=10.2)

    add_docx_section_heading(doc, "Projects")
    for item in RESUME["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(item["heading"]), size=10.7, bold=True)
        add_docx_bullets(doc, item["bullets"])

    add_docx_section_heading(doc, "Professional Experience")
    for item in RESUME["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(item["heading"]), size=10.7, bold=True)
        add_docx_bullets(doc, item["bullets"])

    add_docx_section_heading(doc, "Leadership And Volunteering")
    for item in RESUME["leadership"]:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(item["heading"]), size=10.7, bold=True)
        add_docx_bullets(doc, item["bullets"])

    add_docx_section_heading(doc, "Education")
    p = doc.add_paragraph()
    set_run_font(p.add_run(RESUME["education"]), size=10.4)

    add_docx_section_heading(doc, "Languages")
    p = doc.add_paragraph()
    set_run_font(p.add_run(RESUME["languages"]), size=10.4)

    doc.save(DOCX_PATH)


def wrap_pdf_text(pdf: canvas.Canvas, text: str, width: float, font_name: str, font_size: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if pdf.stringWidth(candidate, font_name, font_size) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(
    pdf: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    width: float,
    font_name: str = "Helvetica",
    font_size: float = 9.4,
    leading: float = 12,
) -> float:
    pdf.setFont(font_name, font_size)
    for line in wrap_pdf_text(pdf, text, width, font_name, font_size):
        pdf.drawString(x, y, line)
        y -= leading
    return y


def draw_contact_lines(pdf: canvas.Canvas, items: Iterable[str], x: float, y: float, width: float) -> float:
    font_name = "Helvetica"
    font_size = 8.3
    separator = " | "
    lines: list[str] = []
    current = ""
    for item in items:
        candidate = item if not current else f"{current}{separator}{item}"
        if pdf.stringWidth(candidate, font_name, font_size) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = item
    if current:
        lines.append(current)

    pdf.setFont(font_name, font_size)
    for line in lines:
        pdf.drawCentredString(x + width / 2, y, line)
        y -= 10.5
    return y


def ensure_pdf_space(pdf: canvas.Canvas, y: float, needed: float, margin: float) -> float:
    if y - needed >= margin:
        return y
    pdf.showPage()
    return LETTER[1] - margin


def section_title(pdf: canvas.Canvas, title: str, x: float, y: float, width: float) -> float:
    pdf.setStrokeColor(colors.HexColor("#2E74B5"))
    pdf.setLineWidth(0.7)
    pdf.line(x, y - 3, x + width, y - 3)
    pdf.setFillColor(colors.HexColor("#2E74B5"))
    pdf.setFont("Helvetica-Bold", 10.5)
    pdf.drawString(x, y + 3, title.upper())
    pdf.setFillColor(colors.black)
    return y - 17


def draw_bullets(pdf: canvas.Canvas, bullets: Iterable[str], x: float, y: float, width: float, margin: float) -> float:
    for bullet in bullets:
        y = ensure_pdf_space(pdf, y, 34, margin)
        lines = wrap_pdf_text(pdf, bullet, width - 14, "Helvetica", 8.9)
        pdf.setFont("Helvetica", 8.9)
        for index, line in enumerate(lines):
            prefix = "- " if index == 0 else "  "
            pdf.drawString(x, y, f"{prefix}{line}")
            y -= 11
        y -= 1
    return y


def build_pdf() -> None:
    pdf = canvas.Canvas(str(PDF_PATH), pagesize=LETTER)
    page_width, page_height = LETTER
    margin = 42
    x = margin
    width = page_width - (margin * 2)
    y = page_height - margin

    pdf.setTitle("Eslam Tarek Nabawy Resume")
    pdf.setAuthor("Eslam Tarek Nabawy")

    pdf.setFont("Helvetica-Bold", 20)
    pdf.drawCentredString(page_width / 2, y, RESUME["name"].upper())
    y -= 20
    pdf.setFont("Helvetica-Bold", 10.5)
    pdf.setFillColor(colors.HexColor("#2E74B5"))
    pdf.drawCentredString(page_width / 2, y, RESUME["title"])
    pdf.setFillColor(colors.black)
    y -= 15
    y = draw_contact_lines(pdf, RESUME["contact"], x, y, width)
    y -= 10

    y = section_title(pdf, "Professional Summary", x, y, width)
    y = draw_wrapped(pdf, RESUME["summary"], x, y, width, font_size=9.2, leading=11.8)
    y -= 7

    y = section_title(pdf, "Technical Skills", x, y, width)
    for label, value in RESUME["skills"]:
        y = ensure_pdf_space(pdf, y, 18, margin)
        pdf.setFont("Helvetica-Bold", 8.9)
        label_text = f"{label}: "
        pdf.drawString(x, y, label_text)
        label_width = pdf.stringWidth(label_text, "Helvetica-Bold", 8.9)
        y = draw_wrapped(pdf, value, x + label_width, y, width - label_width, font_size=8.9, leading=10.8)
    y -= 6

    y = section_title(pdf, "Projects", x, y, width)
    for item in RESUME["projects"]:
        y = ensure_pdf_space(pdf, y, 45, margin)
        pdf.setFont("Helvetica-Bold", 9.4)
        pdf.drawString(x, y, item["heading"])
        y -= 12
        y = draw_bullets(pdf, item["bullets"], x + 8, y, width - 8, margin)
        y -= 2

    y = section_title(pdf, "Professional Experience", x, y, width)
    for item in RESUME["experience"]:
        y = ensure_pdf_space(pdf, y, 45, margin)
        pdf.setFont("Helvetica-Bold", 9.4)
        pdf.drawString(x, y, item["heading"])
        y -= 12
        y = draw_bullets(pdf, item["bullets"], x + 8, y, width - 8, margin)
        y -= 2

    y = section_title(pdf, "Leadership And Volunteering", x, y, width)
    for item in RESUME["leadership"]:
        y = ensure_pdf_space(pdf, y, 45, margin)
        pdf.setFont("Helvetica-Bold", 9.4)
        pdf.drawString(x, y, item["heading"])
        y -= 12
        y = draw_bullets(pdf, item["bullets"], x + 8, y, width - 8, margin)
        y -= 2

    y = ensure_pdf_space(pdf, y, 48, margin)
    y = section_title(pdf, "Education", x, y, width)
    y = draw_wrapped(pdf, RESUME["education"], x, y, width, font_size=9.1, leading=11.5)
    y -= 5

    y = ensure_pdf_space(pdf, y, 38, margin)
    y = section_title(pdf, "Languages", x, y, width)
    draw_wrapped(pdf, RESUME["languages"], x, y, width, font_size=9.1, leading=11.5)
    pdf.save()


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    shutil.copyfile(DOCX_PATH, LEGACY_DOCX_PATH)
    shutil.copyfile(PDF_PATH, LEGACY_PDF_PATH)
    shutil.copyfile(PDF_PATH, PUBLIC_RESUME)
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")
    print(f"Updated {LEGACY_DOCX_PATH}")
    print(f"Updated {LEGACY_PDF_PATH}")
    print(f"Updated {PUBLIC_RESUME}")


if __name__ == "__main__":
    main()
